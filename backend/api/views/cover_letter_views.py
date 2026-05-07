import json
import os
import tempfile
import uuid
from pathlib import Path

from django.conf import settings
from django.http import HttpResponse
from rest_framework import status
from rest_framework.decorators import api_view, parser_classes, permission_classes
from rest_framework.parsers import FormParser, MultiPartParser
from rest_framework.permissions import AllowAny
from rest_framework.response import Response

from .doc_convert_views import MAX_UPLOAD_BYTES
from .tool_chat_static import (
    BEDROCK_ANTHROPIC_VERSION,
    DEFAULT_AWS_REGION,
    DEFAULT_BEDROCK_MODEL_ID,
    ENV_AWS_REGION,
    ENV_BEDROCK_MODEL_ID,
    ERROR_ASSISTANT_REQUEST_FAILED,
    ERROR_BEDROCK_SDK_MISSING,
    ERROR_CHAT_NOT_CONFIGURED,
)

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_JOB_DESCRIPTION_CHARS = 12000
MAX_COVER_LETTER_CHARS = 10000
MAX_OUTPUT_TOKENS = 4096

SYSTEM_PROMPT = """You are an expert career writer. Rewrite cover letters so they are tailored to a specific job description while staying truthful to the candidate's existing experience.

Rules:
- Preserve the candidate's identity, contact details, dates, employers, and facts from the original cover letter.
- Emphasize skills, responsibilities, and keywords that are present in the job description only when they are reasonably supported by the original letter.
- Keep the result polished, specific, and professional.
- Do not invent certifications, degrees, employers, metrics, tools, or years of experience.
- Return only the finished cover letter text with natural paragraph breaks. Do not include explanations, headings like "Updated Cover Letter", markdown, bullets, or placeholders.
"""


def _extract_docx_text(input_path: str) -> str:
    from docx import Document

    document = Document(input_path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    return "\n\n".join(paragraph for paragraph in paragraphs if paragraph)


def _build_user_prompt(cover_letter: str, job_description: str) -> str:
    return (
        "Original cover letter:\n"
        f"{cover_letter[:MAX_COVER_LETTER_CHARS]}\n\n"
        "Job description:\n"
        f"{job_description[:MAX_JOB_DESCRIPTION_CHARS]}\n\n"
        "Rewrite the cover letter so it highlights the candidate's most relevant skills for this job. "
        "Keep it close to the original length unless the original is unusually short."
    )


def _call_cover_letter_model(cover_letter: str, job_description: str) -> tuple[str, str]:
    try:
        import boto3
    except Exception:
        raise RuntimeError(ERROR_BEDROCK_SDK_MISSING)

    region = os.getenv(ENV_AWS_REGION, DEFAULT_AWS_REGION).strip() or DEFAULT_AWS_REGION
    model_id = os.getenv(ENV_BEDROCK_MODEL_ID, DEFAULT_BEDROCK_MODEL_ID).strip() or DEFAULT_BEDROCK_MODEL_ID
    user_prompt = _build_user_prompt(cover_letter, job_description)

    try:
        bedrock = boto3.client("bedrock-runtime", region_name=region)
        body = json.dumps(
            {
                "anthropic_version": BEDROCK_ANTHROPIC_VERSION,
                "max_tokens": MAX_OUTPUT_TOKENS,
                "system": SYSTEM_PROMPT,
                "messages": [{"role": "user", "content": user_prompt}],
            }
        )
        response = bedrock.invoke_model(modelId=model_id, body=body)
        result = json.loads(response["body"].read())
        rewritten = (result.get("content", [{}])[0].get("text", "") or "").strip()
    except Exception as exc:
        error_msg = str(exc)
        if "Could not connect to the endpoint URL" in error_msg or "UnrecognizedClientException" in error_msg:
            raise RuntimeError(ERROR_CHAT_NOT_CONFIGURED) from exc
        if settings.DEBUG:
            raise RuntimeError(f"{ERROR_ASSISTANT_REQUEST_FAILED}: {error_msg}") from exc
        raise RuntimeError(ERROR_ASSISTANT_REQUEST_FAILED) from exc

    if not rewritten:
        raise RuntimeError("No updated cover letter returned. Please try again.")

    return rewritten, model_id


def _replace_document_body(input_path: str, output_path: str, updated_text: str) -> None:
    from docx import Document

    document = Document(input_path)
    style = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            style = paragraph.style
            break

    body = document._body._element
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)

    paragraphs = [part.strip() for part in updated_text.split("\n") if part.strip()]
    for paragraph_text in paragraphs:
        paragraph = document.add_paragraph(paragraph_text)
        if style is not None:
            paragraph.style = style

    document.save(output_path)


@api_view(["POST"])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
def cover_letter_update_view(request):
    doc_file = request.FILES.get("file")
    job_description = (request.data.get("jobDescription") or "").strip()

    if doc_file is None:
        return Response({"error": "No file uploaded. Send a DOCX file as 'file'."}, status=400)

    name_lower = (doc_file.name or "").lower()
    content_type = (doc_file.content_type or "").split(";")[0].strip().lower()
    if content_type != DOCX_MIME_TYPE and not name_lower.endswith(".docx"):
        return Response(
            {"error": "Unsupported file type. Please upload a .docx cover letter."},
            status=415,
        )

    if doc_file.size > MAX_UPLOAD_BYTES:
        return Response(
            {"error": f"File too large. Maximum allowed size is {MAX_UPLOAD_BYTES // 1024 // 1024} MB."},
            status=413,
        )

    if not job_description:
        return Response({"error": "Job description is required."}, status=400)

    if len(job_description) > MAX_JOB_DESCRIPTION_CHARS:
        return Response(
            {"error": f"Job description is too long. Maximum is {MAX_JOB_DESCRIPTION_CHARS} characters."},
            status=413,
        )

    with tempfile.TemporaryDirectory(prefix="torensa_cover_letter_") as tmpdir:
        input_path = os.path.join(tmpdir, f"{uuid.uuid4().hex}.docx")
        output_path = os.path.join(tmpdir, f"{uuid.uuid4().hex}.docx")

        with open(input_path, "wb") as fh:
            for chunk in doc_file.chunks():
                fh.write(chunk)

        try:
            cover_letter_text = _extract_docx_text(input_path)
        except Exception:
            return Response(
                {"error": "Could not read that DOCX file. Please upload a valid Word document."},
                status=400,
            )

        if not cover_letter_text:
            return Response({"error": "The uploaded cover letter does not contain readable text."}, status=400)

        try:
            updated_text, model_id = _call_cover_letter_model(cover_letter_text, job_description)
            _replace_document_body(input_path, output_path, updated_text)
        except RuntimeError as exc:
            error_text = str(exc)
            code = status.HTTP_503_SERVICE_UNAVAILABLE if error_text == ERROR_BEDROCK_SDK_MISSING else status.HTTP_502_BAD_GATEWAY
            return Response({"error": error_text}, status=code)

        with open(output_path, "rb") as fh:
            docx_bytes = fh.read()

    original_stem = Path(doc_file.name or "cover-letter").stem
    output_name = f"{original_stem}-tailored.docx"

    response = HttpResponse(docx_bytes, content_type=DOCX_MIME_TYPE)
    response["Content-Disposition"] = f'attachment; filename="{output_name}"'
    response["Content-Length"] = str(len(docx_bytes))
    response["X-Torensa-Model"] = model_id
    return response
